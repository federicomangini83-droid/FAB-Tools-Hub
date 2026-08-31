import os
import io
import json
import datetime
import pandas as pd
from PIL import Image
import pytesseract
import re

# Word
from docx import Document

# PDF
import pdfplumber
from pdf2image import convert_from_path

start_time = datetime.datetime.now()
print("Ora di inizio:", start_time)

DOCUMENT_FOLDER = r"C:\Users\federico.mangini\Downloads\Document_index"
CSV_OUTPUT_FILE = r"C:\Users\federico.mangini\Downloads\Document_index.csv"
JSON_OUTPUT_FILE = r"C:\Users\federico.mangini\Downloads\Document_index.json"
TXT_OUTPUT_FILE = r"C:\Users\federico.mangini\Downloads\Document_index.txt"

pytesseract.pytesseract.tesseract_cmd = r"C:\Users\federico.mangini\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"


def extract_image_ocr_from_bytes(image_bytes):
    try:
        image = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(image, lang="ita+eng")
        if text.strip():
            return f"\n[IMAGE_OCR]\n{text.strip()}\n[/IMAGE_OCR]\n"
    except:
        return ""
    return ""


def table_to_json_text(columns, rows_data):
    table_json = {"columns": columns, "rows": rows_data}
    return f"[TABLE_JSON]\n{json.dumps(table_json, ensure_ascii=False)}\n[/TABLE_JSON]\n"


# ----------------- WORD -----------------

def parse_word(path):
    doc = Document(path)
    rows = []
    heading_stack = {}
    current_heading = ""
    current_text = []

    for p in doc.paragraphs:

        paragraph_text = ""
        for run in p.runs:
            paragraph_text += run.text
            drawings = run.element.xpath('.//a:blip')
            for blip in drawings:
                rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if rId:
                    image_part = doc.part.related_parts[rId]
                    paragraph_text += extract_image_ocr_from_bytes(image_part.blob)

        text = paragraph_text.strip()
        style = p.style.name.lower()

        if any(h in style for h in ["heading"]):
            if current_text or current_heading:
                rows.append({
                    "Document": os.path.basename(path),
                    "Heading": current_heading,
                    "Text": "\n".join(current_text)
                })
                current_text = []

            match = re.search(r'\d+', style)
            level = int(match.group()) if match else 1
            heading_stack[level] = text
            for k in list(heading_stack.keys()):
                if k > level:
                    del heading_stack[k]
            current_heading = ">///>".join(heading_stack[k] for k in sorted(heading_stack))
        else:
            if text:
                current_text.append(text)

    # Tabelle
    for table in doc.tables:
        columns = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        rows_data = []
        for row in table.rows[1:] if len(table.rows) > 1 else []:
            rows_data.append([cell.text.strip() for cell in row.cells])
        current_text.append(table_to_json_text(columns, rows_data))

    if current_text or current_heading:
        rows.append({
            "Document": os.path.basename(path),
            "Heading": current_heading,
            "Text": "\n".join(current_text)
        })

    return rows


# ----------------- PDF -----------------

def parse_pdf(path):
    rows = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            heading_stack = {}
            current_heading = ""
            current_text = []

            # analizziamo riga per riga
            for char in page.chars:  # char ha: text, size, fontname, etc.
                pass  # placeholder, useremo lines
            
            for line in page.lines:
                pass  # placeholder, useremo line extraction

            # più semplice: estraiamo righe con pdfplumber
            for line in page.extract_text().split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Heading detection: font size o maiuscolo
                # pdfplumber non restituisce facilmente font size per extract_text()
                # come semplificazione: righe tutte maiuscole → Heading1
                if line.isupper() and len(line) > 2:
                    if current_text or current_heading:
                        rows.append({
                            "Document": os.path.basename(path),
                            "Heading": current_heading,
                            "Text": "\n".join(current_text)
                        })
                        current_text = []
                    heading_stack[1] = line
                    current_heading = " > ".join(heading_stack[k] for k in sorted(heading_stack))
                else:
                    current_text.append(line)

            # Tabelle
            page_tables_text = ""
            for table in page.extract_tables():
                if table:
                    columns = table[0]
                    rows_data = table[1:] if len(table) > 1 else []
                    page_tables_text += table_to_json_text(columns, rows_data)

            # immagini
            images_text = ""
            # poppler_path specificato direttamente
            images = convert_from_path(path, first_page=i, last_page=i, poppler_path=r"C:\poppler-25.12.0\Library\bin")
            for img in images:
                images_text += extract_image_ocr_from_bytes(img.tobytes())

            # concat testo finale
            final_text = "\n".join(current_text) + "\n" + images_text + "\n" + page_tables_text
            rows.append({
                "Document": os.path.basename(path),
                "Heading": current_heading or f"Page {i}",
                "Text": final_text.strip()
            })
    return rows


# ----------------- PROCESS DOCUMENTS -----------------

def process_documents():
    all_rows = []
    for file in os.listdir(DOCUMENT_FOLDER):
        path = os.path.join(DOCUMENT_FOLDER, file)
        print("Parsing:", file)
        if file.lower().endswith(".docx"):
            all_rows.extend(parse_word(path))
        elif file.lower().endswith(".pdf"):
            all_rows.extend(parse_pdf(path))
    return all_rows


# ----------------- MAIN -----------------

def main():
    rows = process_documents()

    # CSV
    df = pd.DataFrame(rows)
    df = df[["Document", "Heading", "Text"]]
    df.to_csv(CSV_OUTPUT_FILE, index=False, encoding="utf-8-sig")
    print("CSV creato:", CSV_OUTPUT_FILE)

    # JSON
    with open(JSON_OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(rows, f, indent=2, ensure_ascii=False)
    print("JSON creato:", JSON_OUTPUT_FILE)

    # TXT
    with open(TXT_OUTPUT_FILE, "w", encoding="utf-8") as f_txt:
        for row in rows:
            # ogni riga: Document | Heading | Text
            f_txt.write(f"Document: {row['Document']}\n")
            f_txt.write(f"Heading: {row['Heading']}\n")
            f_txt.write(f"Text:\n{row['Text']}\n")
            f_txt.write("\n---\n\n")  # separatore tra record

    print("TXT creato:", TXT_OUTPUT_FILE)

if __name__ == "__main__":
    main()
    end_time = datetime.datetime.now()
    print("Ora di fine:", end_time)
    print("Durata:", end_time - start_time)
    print("Codice Terminato")